import os
import re
from typing import List, Dict


def extract_text(file_path: str) -> List[Dict]:
    """
    Extract text from a PDF or DOCX file.

    Returns a list of sections:
        [{"title": str, "content": str, "index": int}]

    For PDF: uses PyMuPDF (fitz), detects headings by font size > body average.
    For DOCX: uses python-docx, detects headings by style name containing "Heading".
    If no headings are detected, splits into 3000-word chunks with 200-word overlap.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(file_path: str) -> List[Dict]:
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    all_spans = []

    for page_num, page in enumerate(doc):
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span["text"].strip()
                    if text:
                        all_spans.append({
                            "text": text,
                            "size": span["size"],
                            "page": page_num + 1,
                        })

    doc.close()

    if not all_spans:
        return [{"title": "Document", "content": "", "index": 0}]

    # Compute average font size to detect headings
    sizes = [s["size"] for s in all_spans]
    avg_size = sum(sizes) / len(sizes)
    heading_threshold = avg_size * 1.15  # 15% larger than average = heading

    sections: List[Dict] = []
    current_title = "Introduction"
    current_content_parts: List[str] = []
    section_index = 0

    for span in all_spans:
        if span["size"] >= heading_threshold and len(span["text"]) < 200:
            # Save previous section if it has content
            if current_content_parts:
                sections.append({
                    "title": current_title,
                    "content": " ".join(current_content_parts).strip(),
                    "index": section_index,
                })
                section_index += 1
                current_content_parts = []
            current_title = span["text"]
        else:
            current_content_parts.append(span["text"])

    # Flush last section
    if current_content_parts:
        sections.append({
            "title": current_title,
            "content": " ".join(current_content_parts).strip(),
            "index": section_index,
        })

    if not sections:
        return _fallback_chunk([s["text"] for s in all_spans])

    # If we ended up with only one section and no real headings, fall back to chunking
    if len(sections) == 1:
        return _fallback_chunk([s["text"] for s in all_spans])

    return sections


def _extract_docx(file_path: str) -> List[Dict]:
    from docx import Document

    doc = Document(file_path)
    sections: List[Dict] = []
    current_title = "Introduction"
    current_content_parts: List[str] = []
    section_index = 0
    found_headings = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        is_heading = "heading" in style_name.lower()

        if is_heading:
            found_headings = True
            if current_content_parts:
                sections.append({
                    "title": current_title,
                    "content": " ".join(current_content_parts).strip(),
                    "index": section_index,
                })
                section_index += 1
                current_content_parts = []
            current_title = text
        else:
            current_content_parts.append(text)

    # Flush last section
    if current_content_parts:
        sections.append({
            "title": current_title,
            "content": " ".join(current_content_parts).strip(),
            "index": section_index,
        })

    if not found_headings or len(sections) <= 1:
        all_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return _fallback_chunk(all_text)

    return sections


def _fallback_chunk(text_parts: List[str], chunk_words: int = 3000, overlap_words: int = 200) -> List[Dict]:
    """Split flat text into overlapping word chunks."""
    full_text = " ".join(text_parts)
    words = full_text.split()

    if not words:
        return [{"title": "Document", "content": "", "index": 0}]

    sections = []
    start = 0
    index = 0

    while start < len(words):
        end = min(start + chunk_words, len(words))
        chunk_text = " ".join(words[start:end])
        sections.append({
            "title": f"Section {index + 1}",
            "content": chunk_text,
            "index": index,
        })
        if end == len(words):
            break
        start = end - overlap_words
        index += 1

    return sections
